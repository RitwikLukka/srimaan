# -*- coding: utf-8 -*-
from __future__ import unicode_literals
from django.shortcuts import render,redirect
from django.http import HttpResponse
from django.shortcuts import redirect
from django.template import loader
from django.forms import modelformset_factory
from .models import HomeR1C1,HomeR1C2,HomeR1C3A,HomeR2,HomeR3,HomeR4,Gurukulam,Kainkariyam,Library,MappingTable,Donar,Archives,ArchiveImg,Eventdetails,Adminupload
from docx import Document
from django.template.loader import render_to_string
from django.http import HttpResponseRedirect,JsonResponse,HttpResponse
from django.db.models import Q
import os
import pypandoc
from django.conf import settings
from django.core.urlresolvers import reverse
#from django.core.urlresolvers import reverse
from paypal.standard.forms import PayPalPaymentsForm
from django.views.decorators.csrf import csrf_exempt
from .models import Question,Reference,Adminupload,Test,Tt,Demo,Subs,Jobs
from .forms import question,upload,uploads,upld,r,r3,subs
from .forms import LoginForm
from .models import Feedback
from .forms import feedback,ImageForm
from .forms import validation
from .models import Admin,Images
from .models import Courses
from .models import Job
from .models import Job1
from .forms import jb
from .forms import cour
import datetime,time


# Create your views here.
def index(request):
    r1c1 = HomeR1C1.objects.all()
    r1c2 = HomeR1C2.objects.all()
    r1c3a = HomeR1C3A.objects.all()
    r2 = HomeR2.objects.all()
    r5 = HomeR4.objects.all()
    r3 = HomeR3.objects.filter(active=1).order_by('-updated_on')[0:3]
    data=Feedback.objects.all().order_by('-id')
    filename = Archives.objects.filter(active=1).order_by('-updated_on')[:3]
    print(filename)
    size=len(filename)
    r4=[]
    for i in range(0,size):
        obj=ArchiveImg.objects.filter(Q(name=filename[i].title) & Q(updated_on=filename[i].updated_on))[:1]
        r4+=obj


    return render(request, 'home/index.html', {'r1c1':r1c1,'r1c2':r1c2,'r1c3a':r1c3a,'r2':r2,'r3':r3,'r4':r4,'data':data,'r5':r5})
def archivesearch(request):
    if request.method=='POST':
        key = request.POST.get('key','')
        details=Archives.objects.filter(title=key)
        size=len(details)
        if (size < 1):
            print('enter the ifs')
            print(size)
            return render(request,'home/archive2.html',{'size':size})

        r4=[]
        for i in range(0,size):
            obj=ArchiveImg.objects.filter(Q(name=details[i].title) & Q(updated_on=details[i].updated_on))
            r4+=[obj]

        g=Archives.objects.filter(title=key)
        document = Document()
        for i in range(0,size):
            file=str(details[i].content)
            path=os.path.join(settings.MEDIA_ROOT,file)
            output=pypandoc.convert(path,'html')
            document.add_paragraph(output)
        file_content=document.paragraphs
        return render(request, 'home/archive2.html',{'file_content':file_content,'r4':r4,'g':g,'key':key})

    else:
        filename = Archives.objects.filter(active=1).order_by('-updated_on')[:5]
        size=len(filename)

        r4=[]
        for i in range(0,size):
            obj=ArchiveImg.objects.filter(Q(name=filename[i].title) & Q(updated_on=filename[i].updated_on))
            r4+=[obj]
        print(r4)
        g=Archives.objects.filter()
        document = Document()
        for i in range(0,size):
            file=str(filename[i].content)
            path=os.path.join(settings.MEDIA_ROOT,file)
            output=pypandoc.convert(path,'html')
            document.add_paragraph(output)
        file_content=document.paragraphs
        return render(request, 'home/archive1.html',{'file_content':file_content,'r4':r4,'g':g})
def archives1(request):
    saved = False
    ts=time.time()
    timestamp=datetime.datetime.fromtimestamp(ts).strftime('%Y-%m-%d')
    if request.method == "POST":

        profile = Archives()

        profile.title = request.POST.get('content','')
        profile.image = request.POST.get('image','')
        profile.content = request.POST.get('link','')
        profile.updated_by = "admin"
        profile.updated_on = timestamp
        profile.active = 1
        profile.save()
        saved = True
        return render(request,'home/save1.html',locals())
    else:

        saved = 'false'
        return render(request,'home/save1.html',{'saved':saved})

def archives1img(request):
    saved = False
    ts=time.time()
    timestamp=datetime.datetime.fromtimestamp(ts).strftime('%Y-%m-%d')
    if request.method == "POST":

        profile = ArchivesImg()
        profile.name = request.POST.get('content','')
        profile.img = request.POST.get('image','')
        profile.updated_by = "admin"
        profile.updated_on = timestamp
        profile.active = 1
        profile.save()
        saved = True
        return render(request,'home/save1.html',locals())
    else:

        saved = 'false'
        return render(request,'home/save1.html',{'saved':saved})
def donate(request):
    return render(request, 'home/donate.html')
    #donate_url = loader.get_template('home/donate.html')
    # return HttpResponse(donate_url.render(request))
def sub(request):
    return render(request, 'home/sub.html')

def fb(request):
    saved = False

    if request.method == "POST":
        My = feedback(request.POST)
        if My.is_valid():
            profile = Feedback()
            profile.name = My.cleaned_data["name"]
            profile.city = My.cleaned_data["city"]
            profile.feedback = My.cleaned_data["feedback"]


            profile.save()
            saved = True
        else:
            My = feedback()
        return render(request,'home/save.html',locals())
def log(request):
    data= Question.objects.all()
    return render(request, 'home/q&a.html',{"data": data})

def admin(request):
    return render(request, 'home/login.html')

def red(request):
    return Redirect("www.srimaaninfotech.org")

def than(request):
    return render(request, "home/test1.html")
def thanks(request):
    r1c1 = Adminupload.objects.all()
    return render(request, "home/test.html", {'r1c1': r1c1})

def about_trust(request):
    return render(request, 'home/about_trust.html')

def about_bhattar(request):
    return render(request, 'home/about_bhattar.html')

def gurukulam(request):
    filename = (Gurukulam.objects.filter(active=1).order_by('-updated_on')[0]).content
    r4 = HomeR4.objects.all()
    file=str(filename)
    path=os.path.join(settings.MEDIA_ROOT,file)
    output=pypandoc.convert(path,'html')
    document = Document()
    document.add_paragraph(output)
    file_content=document.paragraphs
    return render(request, 'home/gurukulam.html',{'file_content':file_content,'r4':r4})

def kainkariyam(request):
    filename = (Kainkariyam.objects.filter(active=1).order_by('-updated_on')[0]).content
    document = Document(filename)
    file_content = document.paragraphs
    # file_open.close()
    return render(request, 'home/kainkariyam.html',{'file_content':file_content})

def library(request):
    filename = (Library.objects.filter(active=1).order_by('-updated_on')[0]).content
    document = Document(filename)
    file_content = document.paragraphs
    # file_open.close()
    return render(request, 'home/library.html',{'file_content':file_content})

def payment_process(request):
    host=request.get_host()
    # try:
    #     amount=request.session['amount']
    # except:
    #     pass
    paypal_dict={
    'business':settings.PAYPAL_RECEIVER_EMAIL,
    'amount':request.session['amount'],
    'currency_code':'INR',
    'notify_url':'http://{}{}'.format(host,reverse('paypal-ipn')),
    'return_url':'http://{}{}'.format(host,reverse('payment_done')),
    'cancel_return':'http://{}{}'.format(host,reverse('payment_cancelled'))}
    #form=PayPalPaymentsForm(initial=paypal_dict)
    return render(request,'home/process.html',{'paypal_dict':paypal_dict})

@csrf_exempt
def payment_done(request):
    return render(request,'home/done.html')

@csrf_exempt
def payment_cancelled(request):
    if request.session.has_key('name'):
        try:
            del request.session['name']
            del request.session['amount']
            return render(request,'home/cancelled.html')
        except:
            pass
    else:
        return redirect('home.views.checkout')

def page(request,name):
    user=Reference.objects.get(id=name)
    print(user)
    nam=user.name
    filename = (Adminupload.objects.filter(reference_id=name).order_by('-updated_on')[0]).link
    print(filename)
    r4 = (Images.objects.filter(reference_id=name).order_by('-updated_on')[:19])
    print(r4)
    file=str(filename)
    print(file)
    path=os.path.join(settings.MEDIA_ROOT,file)
    output=pypandoc.convert(path,'html')
    document = Document()
    document.add_paragraph(output)
    file_content=document.paragraphs
    return render(request, 'home/page.html',{'file_content':file_content,'r4':r4,'name':nam})


def donar(request):
    name = Donar.objects.get(id=1)
    filename=name.link
    print(filename)
    r4 = HomeR4.objects.all()
    file=str(filename)
    path=os.path.join(settings.MEDIA_ROOT,file)
    print(path)
    output=pypandoc.convert(path,'html')
    document = Document()
    document.add_paragraph(output)
    file_content=document.paragraphs
    print(file)
    return render(request, 'home/donar.html',{'file_content':file_content})


def checkout(request,name):
    user=MappingTable.objects.get(id=name)
    nam=user.donation_items
    if request.method=='POST':
        firstname=request.POST.get('fname','')
        lastname=request.POST.get('lname','')
        emailid=request.POST.get('EmailId','')
        amount=request.POST.get('amount','')
        number=request.POST.get('contact','')
        request.session['amount']=request.POST.get('amount','')
        request.session['name']=firstname

        return redirect('home.views.payment_process')
        # firstname=request.POST.get('fname','')
        # lastname=request.POST.get('lname','')
        #obj=details(username=username,password=password)
        #obj.save()
    else:
         return render(request,'home/checkout.html',{'name':nam})

def SaveProfile(request):
    saved = False

    if request.method == "POST":
        MyProfileForm = question(request.POST)
        if MyProfileForm.is_valid():
            profile = Question()
            profile.name = MyProfileForm.cleaned_data["Name"]
            profile.emailid = MyProfileForm.cleaned_data["EmailId"]
            profile.phone = MyProfileForm.cleaned_data["phone"]
            profile.question = MyProfileForm.cleaned_data["question"]


            profile.save()
            saved = True
        else:
            MyProfileForm = question()
        return render(request,'home/save.html',locals())

def login(request):
    if request.is_ajax():

        flag=request.GET.get('num', None)
        ImageFormSet = modelformset_factory(Images,form=ImageForm, extra=int(flag))
        formset = ImageFormSet(queryset=Adminupload.objects.none())
        # data1=json.dumps(formset.__dict__)
        data={
        'formset':formset
        }
        html = render_to_string('home/ajax_img.html',data,request=request)
        return JsonResponse(html,safe=False)



    username = "not logged in"
    # if request.session.has_key('username'):
    #     username = request.session['username']
    #     return render(request,'home/adminlg.html',{"username":username})
    if request.method == "POST":
        #Get the posted form
        MyLoginForm = validation(request.POST)
        if MyLoginForm.is_valid():
            username = MyLoginForm.clean_message()
            print(username)
            #request.session['username'] = username
        else:
            MyLoginForm = validation()

    if username == "not logged in":
        return render(request,'home/login.html')
    if(username =="login"):
        return render(request,'home/login.html')
    else:
        form = upload()
        ra=r()
        ImageFormSet = modelformset_factory(Images,form=ImageForm, extra=0)
        formset = ImageFormSet(queryset=Images.objects.none())
        return render(request, 'home/adminlg.html',{'form':form,'formset':formset,'ra':ra})

def logi(request):
    form = upload()
    return render(request, 'home/adminlg.html',{'form':form})

def adminup(request):




    # saved = False

    # if request.method == "POST":
    #
    #     MyProfile = uploads(request.POST,request.FILES)
    #     if MyProfile.is_valid():
    #
    #         profile = Test()
    #         profile.content = MyProfile.cleaned_data["content"]
    #         profile.image = MyProfile.cleaned_data["image"]
    #         profile.link = MyProfile.cleaned_data["link"]
    #
    #
    #         profile.save()
    #         saved = True
    #     else:
            MyProfile = upload()
            return render(request,'home/save.html',locals())

def adup(request,name):

    saved = False

    ts=time.time()
    timestamp=datetime.datetime.fromtimestamp(ts).strftime('%Y-%m-%d %H:%M:%S')
    if request.method == "POST":

        MyProfile = upload(request.POST,request.FILES)
        num=request.POST.get('No_of_pics','')
        ImageFormSet = modelformset_factory(Images,form=ImageForm, extra=num)
        formset = ImageFormSet(request.POST, request.FILES,queryset=Images.objects.none())
        if MyProfile.is_valid() and formset.is_valid():
            profile = Adminupload()
            img     = Images()
            profile.reference_id = name
            profile.content = request.POST.get("content")
            profile.link = MyProfile.cleaned_data["link"]
            profile.updated_by = "admin"
            profile.updated_on = timestamp
            profile.active = 1
            profile.save()
            for form in formset.cleaned_data:
            #             #this helps to not crash if the user
            #             #do not upload all the photos
                if form:

                   image = form['image']
                   photo = Images(reference_id=name,image=image,updated_on = timestamp)
                   photo.save()

            saved = True
            return render(request,'home/save1.html',locals())
        else:
            print(MyProfile.errors and formset.errors)


    else:
        form = upload()
        return render(request,'home/adminlg.html',{'form':form})

def tests(request):
    saved = False

    ts=time.time()
    timestamp=datetime.datetime.fromtimestamp(ts).strftime('%Y-%m-%d %H:%M:%S')
    if request.method == "POST":
        MyProfile = r(request.POST,request.FILES)

        if MyProfile.is_valid():

            profile = HomeR4()
            profile.year = MyProfile.cleaned_data["year"]
            profile.image = MyProfile.cleaned_data["image"]
            profile.content = MyProfile.cleaned_data["content"]
            profile.updated_by = "Admin"
            profile.updated_on = timestamp
            profile.active = 1
            profile.save()
            saved = True
            return render(request,'home/save1.html',locals())

        else:
            print(MyProfile.errors)
    else:
        MyProfile = r()
        return render(request,'home/save1.html',locals())

def archives(request):
    if request.method=='POST':
        from_date=request.POST.get('from','')
        to_date=request.POST.get('to','')
        filename=Archives.objects.filter(updated_on__range=[from_date, to_date])
        size=len(filename)
        print(filename)
        r4=[]
        for i in range(0,size):
            obj=ArchiveImg.objects.filter(Q(name=filename[i].title) & Q(updated_on=filename[i].updated_on))
            r4+=[obj]
        print(r4)
        # g=Archives.objects.filter()
        document = Document()
        for i in range(0,size):
            file=str(filename[i].content)
            path=os.path.join(settings.MEDIA_ROOT,file)
            output=pypandoc.convert(path,'html')
            document.add_paragraph(output)
        file_content=document.paragraphs
        return render(request, 'home/archive.html',{'file_content':file_content,'r4':r4,'filename':filename})
    else:
        filename = Archives.objects.filter(active=1).order_by('-updated_on')[:5]
        size=len(filename)

        r4=[]
        for i in range(0,size):
            obj=ArchiveImg.objects.filter(Q(name=filename[i].title) & Q(updated_on=filename[i].updated_on))
            r4+=[obj]
        print(r4)
        g=Archives.objects.filter()
        document = Document()
        for i in range(0,size):
            file=str(filename[i].content)
            path=os.path.join(settings.MEDIA_ROOT,file)
            #print(path)
            output=pypandoc.convert(path,'html')
            document.add_paragraph(output)
        file_content=document.paragraphs
        return render(request, 'home/archive1.html',{'file_content':file_content,'r4':r4,'g':g})

def link(request):
    saved = False

    ts=time.time()
    timestamp=datetime.datetime.fromtimestamp(ts).strftime('%Y-%m-%d %H:%M:%S')
    if request.method == "POST":
        MyProfile = r3(request.POST)

        if MyProfile.is_valid():

            profile = HomeR3()
            profile.content = MyProfile.cleaned_data["content"]
            profile.links = MyProfile.cleaned_data["links"]
            profile.updated_by = "Admin"
            profile.updated_on = timestamp
            profile.active = 1
            profile.save()
            saved = True
            return render(request,'home/save1.html',locals())

        else:
            print(MyProfile.errors)
    else:
        MyProfile = r3()
        return render(request,'home/save1.html',locals())

def courses(request):
    saved = False

    if request.method == "POST":
        My = cour(request.POST)
        if My.is_valid():
            profile = Courses()
            profile.firstname = My.cleaned_data["firstname"]
            profile.lastname = My.cleaned_data["lastname"]
            profile.gender = My.cleaned_data["gender"]
            profile.email_address = My.cleaned_data["email_address"]
            profile.education_qualification = My.cleaned_data["education_qualification"]
            profile.course = My.cleaned_data["course"]
            profile.contact_number = My.cleaned_data["contact_number"]
            profile.residential_address = My.cleaned_data["residential_address"]
            profile.save()
            saved = True
        else:
            My = cour()
        return render(request,'home/save1.html',locals())



def jbs(request):
    saved = False

    if request.method == "POST":
        My = jb(request.POST)
        if My.is_valid():
            profile = Job()
            profile.firstname = My.cleaned_data["firstname"]
            profile.lastname = My.cleaned_data["lastname"]
            profile.gender = My.cleaned_data["gender"]
            profile.emailid = My.cleaned_data["emailid"]
            profile.education_qualification = My.cleaned_data["education_qualification"]
            profile.contactnumber = My.cleaned_data["contactnumber"]
            profile.postapplied= My.cleaned_data["postapplied"]
            profile.upload= My.cleaned_data["upload"]
            profile.save()
            saved = True
        else:
            My = jb()
        return render(request,'home/save1.html',locals())


def jbs1(request):
    filename = (Job1.objects.get()).content
    g=Job1.objects.get()
    file=str(filename)
    path=os.path.join(settings.MEDIA_ROOT,file)
    output=pypandoc.convert(path,'html')
    document = Document()
    document.add_paragraph(output)
    file_content=document.paragraphs
    return render(request, 'home/job.html',{'file_content':file_content,'g':g})

def crime(request):
    saved = False
    if request.method == "POST":
        My = subs(request.POST)
        if My.is_valid():
            profile = Subs()
            profile.name = My.cleaned_data["name"]
            profile.phone = My.cleaned_data["phone"]
            profile.email = My.cleaned_data["email"]
            profile.address = My.cleaned_data["address"]
            profile.send_via = My.cleaned_data["send_via"]
            profile.save()
            saved = True
        else:
            print(My.errors)
    else:
        My = subs()
    return render(request,'home/save1.html',locals())

def events(request,id):
    filename = (HomeR1C2.objects.filter(id=id).order_by('-updated_on')[0]).details
    print(filename)
    file=str(filename)
    print(file)
    path=os.path.join(settings.MEDIA_ROOT,file)
    output=pypandoc.convert(path,'html')
    document = Document()
    document.add_paragraph(output)
    file_content=document.paragraphs

    if request.method == "POST":
        print("hi")
        profile = Eventdetails()
        profile.firstname =request.POST.get('firstname','')
        profile.lastname = request.POST.get('lastname','')
        profile.gender = request.POST.get('gender','')
        profile.emailid = request.POST.get('emailid','')
        profile.contactno = request.POST.get('contactnumber','')
        profile.event = request.POST.get('event','')
        profile.save()
        saved = True
        return redirect('home.views.index')
    else:
        return render(request,'home/events.html',{'file_content':file_content})


def Update(request):
    udata = Courses.objects.all()
    return render(request,'home/dboc.html',{"udata":udata})
#uploaded data editing
def Updatedata(request):
    udata = Adminupload.objects.all()
    data = Adminupload.objects.values('reference_id')
    print(data)
    # data1=data.reference_id
    # pdata = Reference.objects.filter(id=data)
    # print(pdata)
    return render(request,'home/uploadeddata.html',{"udata":udata})

#for deleting data
def delete(request):
    name = request.POST['row1']
    #print(name)
    #id = request.GET.get('name')
    #print (id)
    ddata = Adminupload.objects.filter(id=name).delete()
    #print (ddata)
    udata = Adminupload.objects.order_by('-id')
    return render(request,'home/uploadeddata.html',{"udata":udata})

#for editing data
def edit(request):
    id = request.POST['row4']
    print(id)
    edata = Adminupload.objects.get(id=id)
    edata.content = request.POST['con']
    edata.image = request.POST['img']
    edata.link = request.POST['lin']
    #print(edata)
    #print(edata.content)
    edata.save()
    udata = Adminupload.objects.order_by('-id')
    return render(request,'home/uploadeddata.html',{"udata":udata})
def logout(request):
   try:
      del request.session['username']
   except:
      pass
   return render(request,'home/login.html')
def archupload(request):

    saved = False

    ts=time.time()
    timestamp=datetime.datetime.fromtimestamp(ts).strftime('%Y-%m-%d')
    name1 = request.POST.get("content")
    if request.method == "POST":

        MyProfile = upload(request.POST,request.FILES)
        num=request.POST.get('No_of_pics','')
        ImageFormSet = modelformset_factory(Images,form=ImageForm, extra=num)
        formset = ImageFormSet(request.POST, request.FILES,queryset=Images.objects.none())
        if MyProfile.is_valid() and formset.is_valid():
            profile = Archives()
            img     = ArchiveImg()
            profile.title = request.POST.get("content")
            profile.content = MyProfile.cleaned_data["link"]
            profile.updated_by = "admin"
            profile.updated_on = timestamp
            profile.active = 1
            profile.save()
            for form in formset.cleaned_data:
            #             #this helps to not crash if the user
            #             #do not upload all the photos
                if form:

                   image = form['image']
                   print(name1)
                   photo = ArchiveImg(name=name1,img=image,updated_by = Admin,updated_on = timestamp,active ="1")
                   photo.save()

            saved = True
            return render(request,'home/save1.html',locals())
        else:
            print (MyProfile.errors and formset.errors)


    else:
        form = upload()
        return render(request,'home/adminlg.html',{'form':form})

def jobupload(request):
    saved = False
    ts=time.time()
    timestamp=datetime.datetime.fromtimestamp(ts).strftime('%Y-%m-%d %H:%M:%S')
    if request.method == "POST":
        profile = Jobs()
        profile.document = request.POST.get("link")
        profile.updated_by = "Admin"
        profile.updated_on = timestamp
        profile.active = 1
        profile.save()
        saved = True
    else:
        My = subs()
    return render(request,'home/save1.html',locals())
